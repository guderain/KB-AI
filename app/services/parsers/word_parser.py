import logging
from pathlib import Path

from docx import Document as DocxDocument
from langchain_core.documents import Document

logger = logging.getLogger(__name__)


class WordParser:
    def parse(self, path: Path) -> list[Document]:
        docs: list[Document] = []
        try:
            doc = DocxDocument(str(path))
        except Exception as exc:
            logger.warning('Failed to open Word %s: %s', path, exc)
            return docs

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            docs.append(
                Document(
                    page_content=text,
                    metadata={
                        'source': str(path),
                        'doc_type': 'word',
                        'element_type': 'text',
                    },
                )
            )

        for table in doc.tables:
            rows_text: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_text.append(' | '.join(cells))
            if rows_text:
                table_text = '\n'.join(rows_text)
                docs.append(
                    Document(
                        page_content=table_text,
                        metadata={
                            'source': str(path),
                            'doc_type': 'word',
                            'element_type': 'table',
                        },
                    )
                )

        return docs
